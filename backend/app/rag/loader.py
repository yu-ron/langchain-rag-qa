"""
文档加载器
支持多种文件格式：PDF、TXT、CSV、DOCX、Markdown
把不同类型的文件统一转换成 LangChain 的 Document 对象
"""
import os
from typing import List
from langchain_core.documents import Document


async def load_document(file_path: str, filename: str) -> List[Document]:
    """
    根据文件类型选择合适的加载器，加载文档内容
    返回 LangChain Document 对象列表
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return _load_pdf(file_path)
    elif ext == ".txt":
        return _load_txt(file_path)
    elif ext == ".csv":
        return _load_csv(file_path)
    elif ext in (".md", ".markdown"):
        return _load_txt(file_path)  # Markdown 本质是纯文本
    elif ext in (".docx", ".doc"):
        return _load_docx(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {ext}")


def _load_pdf(file_path: str) -> List[Document]:
    """
    加载 PDF 文件
    优先用 pdfplumber（解析质量好），出问题用 pypdf 兜底
    """
    try:
        import pdfplumber
        docs = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    docs.append(Document(
                        page_content=text.strip(),
                        metadata={"page": i + 1, "source": file_path},
                    ))
        if docs:
            return docs
    except Exception:
        pass  # pdfplumber 失败，降级到 pypdf

    # 兜底：使用 pypdf
    from langchain_community.document_loaders import PyPDFLoader
    loader = PyPDFLoader(file_path)
    return loader.load()


def _load_txt(file_path: str) -> List[Document]:
    """
    加载纯文本 / Markdown 文件
    不同编码尝试读取
    """
    for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read()
            if text.strip():
                return [Document(
                    page_content=text,
                    metadata={"source": file_path},
                )]
        except UnicodeDecodeError:
            continue
    raise ValueError("无法识别文件编码")


def _load_csv(file_path: str) -> List[Document]:
    """
    加载 CSV 文件
    每行转成一段文字描述（适合商品信息表格）
    """
    import csv
    docs = []
    for encoding in ["utf-8", "gbk", "gb2312"]:
        try:
            with open(file_path, "r", encoding=encoding, newline="") as f:
                reader = csv.DictReader(f)
                for row_num, row in enumerate(reader, start=2):
                    # 把一行 CSV 转成可读的文字
                    parts = [f"{k}: {v}" for k, v in row.items() if v]
                    content = " | ".join(parts)
                    if content.strip():
                        docs.append(Document(
                            page_content=content,
                            metadata={"row": row_num, "source": file_path},
                        ))
            if docs:
                return docs
        except (UnicodeDecodeError, csv.Error):
            continue
    return [Document(page_content="", metadata={"source": file_path})]


def _load_docx(file_path: str) -> List[Document]:
    """
    加载 Word 文档 (.docx)
    """
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        full_text = "\n".join([
            para.text for para in doc.paragraphs if para.text.strip()
        ])
        if full_text.strip():
            return [Document(
                page_content=full_text,
                metadata={"source": file_path},
            )]
    except Exception as e:
        raise ValueError(f"Word 文档解析失败: {e}")

    return [Document(page_content="", metadata={"source": file_path})]
